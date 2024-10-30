#!/usr/bin/env python3
"""
Extract paragraphs with specific styles from a Word document.
"""

import logging
from pathlib import Path
import argparse
import sys
from docx import Document

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def show_usage():
    """
    Display usage instructions for the extract_headings.py script.

    This function prints a formatted message to the console, providing
    information on how to use the script, including command-line syntax
    for listing styles and extracting specific heading styles from a
    Word document.

    The usage instructions cover two main scenarios:
    1. Listing all styles in a document
    2. Extracting paragraphs with specific heading styles

    Examples of command usage are also provided.

    Parameters:
    None

    Returns:
    None

    Note:
    This function is typically called when the script is run without
    arguments or when the user needs help understanding how to use
    the script.
    """
    print("""
Usage Instructions:
------------------
1. To list all styles in a document:
   python extract_headings.py <document_name.docx> --list-styles

2. To extract specific heading styles:
   python extract_headings.py <document_name.docx> --styles ".Head 1,.Head 2,.Head 3"

Examples:
---------
python extract_headings.py mydoc.docx --list-styles
python extract_headings.py mydoc.docx --styles ".Head 1,.Head 2"

Note: The document name is required. If you don't provide it, this help message will be shown.
    """)


def list_all_styles(doc_path: str) -> None:
    """
    List all unique style names in the document along with a sample of text for each style.

    This function opens a Word document, iterates through its paragraphs, and prints
    out each unique style name encountered along with a sample of text using that style.
    It also provides a count of the total number of unique styles found in the document.

    Args:
        doc_path (str): The file path to the Word document to be processed.

    Returns:
        None: This function doesn't return a value, it prints the results to the console.

    Note:
        The function prints the style names and text samples to the console.
        For each style, it shows up to 50 characters of sample text.
    """
    doc = Document(doc_path)
    styles = set()

    print("\nAll Styles in Document:")
    print("-" * 50)

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name not in styles:
            styles.add(style_name)
            # Show a sample of text using this style
            text_sample = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
            print(f"\nStyle: {style_name}")
            print(f"Sample text: {text_sample}")

    print(f"\nTotal unique styles found: {len(styles)}")


def extract_styled_paragraphs(doc_path: str, target_styles: list) -> list:
    """
    Extract paragraphs from a Word document that match the specified styles.

    This function opens a Word document, iterates through its paragraphs,
    and extracts those whose styles match the provided target styles.

    Args:
        doc_path (str): The file path to the Word document to be processed.
        target_styles (list): A list of style names to match against. 
                              The matching is case-insensitive.

    Returns:
        list: A list of tuples, where each tuple contains two elements:
              (style_name: str, text: str). The style_name is the name of 
              the paragraph's style, and text is the content of the paragraph.

    Raises:
        Exception: If there's an error in processing the Word document, 
                   the function logs the error and re-raises the exception.
    """
    try:
        doc = Document(doc_path)
        extracted = []
        target_styles = [style.lower() for style in target_styles]  # Case-insensitive matching

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            text = paragraph.text.strip()

            if not text:
                continue

            if style_name.lower() in target_styles:
                extracted.append((style_name, text))
                logger.info(f"Found matching style '{style_name}': {text[:50]}...")

        return extracted

    except Exception as e:
        logger.error(f"Error processing Word document: {str(e)}")
        raise


def write_to_markdown(paragraphs: list, output_path: str) -> None:
    """
    Write the extracted paragraphs to a markdown file.

    This function takes a list of paragraphs (each containing a style name and text content)
    and writes them to a markdown file. It groups paragraphs by their style, creating section
    headers for each unique style encountered.

    Args:
        paragraphs (list): A list of tuples, where each tuple contains two elements:
                           (style_name: str, text: str). The style_name represents the
                           paragraph's style, and text is the content of the paragraph.
        output_path (str): The file path where the markdown content will be written.

    Returns:
        None

    Raises:
        Exception: If there's an error during the file writing process, the function
                   logs the error and re-raises the exception.

    Note:
        The function uses a logger to record the successful saving of content and any errors.
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            current_style = None

            for style_name, text in paragraphs:
                if style_name != current_style:
                    f.write(f"\n## {style_name}\n\n")
                    current_style = style_name

                f.write(f"{text}\n\n")

        logger.info(f"Content saved to {output_path}")

    except Exception as e:
        logger.error(f"Error writing markdown file: {str(e)}")
        raise


def main():
    """
    Main function to extract styled paragraphs from Word document.

    Parameters:
    input_file (str): Path to the input Word document. If not provided, the function will show usage and exit.
    -s, --styles (str): Comma-separated list of style names to extract (e.g., ".Head 1,.Head 2").
    -o, --output (str): Path for the output markdown file (default: extracted_content.md).
    --list-styles (bool): If True, the function will list all styles in the document and exit.

    Returns:
    int: 0 if the function completes successfully, 1 if an error occurs.
    """
    # If no arguments provided, show usage and exit
    if len(sys.argv) == 1:
        show_usage()
        return 1

    parser = argparse.ArgumentParser(
        description='Extract paragraphs with specific styles from a Word document'
    )
    parser.add_argument(
        'input_file', 
        nargs='?',  # Makes the argument optional
        help='Path to the input Word document'
    )
    parser.add_argument(
        '-s', '--styles',
        help='Comma-separated list of style names to extract (e.g., ".Head 1,.Head 2")'
    )
    parser.add_argument(
        '-o', '--output',
        help='Path for the output markdown file (default: extracted_content.md)',
        default='extracted_content.md'
    )
    parser.add_argument(
        '--list-styles',
        action='store_true',
        help='List all styles in the document and exit'
    )

    args = parser.parse_args()

    # If no input file provided, show usage
    if not args.input_file:
        show_usage()
        return 1

    try:
        input_path = Path(args.input_file)

        if not input_path.exists():
            logger.error(f"Error: Document '{args.input_file}' not found!")
            show_usage()
            return 1

        logger.info(f"Processing Word document: {input_path}")

        # List styles mode
        if args.list_styles:
            list_all_styles(str(input_path))
            return 0

        # Check for styles parameter
        if not args.styles:
            logger.error("No styles specified. Use --styles or --list-styles to see available styles")
            show_usage()
            return 1

        # Extract the specified styles
        target_styles = [s.strip() for s in args.styles.split(',')]
        paragraphs = extract_styled_paragraphs(str(input_path), target_styles)

        if not paragraphs:
            logger.warning(f"No paragraphs found with styles: {', '.join(target_styles)}")
            return 1

        # Print found paragraphs to console
        print("\nExtracted Content:")
        print("-" * 20)
        current_style = None
        for style_name, text in paragraphs:
            if style_name != current_style:
                print(f"\n{style_name}:")
                current_style = style_name
            print(f"  {text[:50]}...")

        # Write to markdown file
        write_to_markdown(paragraphs, args.output)

        return 0

    except Exception as e:
        logger.error(f"An error occurred: {str(e)}")
        show_usage()
        return 1


if __name__ == "__main__":
    exit(main())